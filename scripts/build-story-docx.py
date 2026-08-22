from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "第七次重启_V1.0_主流程_8月4日人工重力整合版.txt"
OUTPUT_NAME = "第七次重启_V1.0_主流程_8月4日人工重力整合版.docx"
OUTPUT = ROOT / OUTPUT_NAME
DESKTOP_OUTPUT = ROOT.parent / OUTPUT_NAME

FONT_CJK = "Microsoft YaHei"
FONT_MONO = "Consolas"
INK = RGBColor(25, 35, 48)
MUTED = RGBColor(99, 115, 132)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(25, 126, 132)
GOLD = RGBColor(122, 90, 0)
SOFT_BLUE = "E8EEF5"
SOFT_GRAY = "F4F6F9"


def set_run_font(run, name: str, size: float, color: RGBColor | None = None, bold: bool = False, italic: bool = False):
    run.font.name = name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_bar(paragraph, color: str = "2E74B5", size: str = "12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), color)
    borders.append(left)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, FONT_MONO, 8, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_section(section):
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("SEVENTH REBOOT  /  STORY MASTER  /  V1.0")
    set_run_font(hr, FONT_MONO, 7.5, MUTED, True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = fp.add_run("AUG 04  /  ")
    set_run_font(label, FONT_MONO, 7.5, MUTED)
    add_page_field(fp)


def configure_style(style, font: str, size: float, color: RGBColor, before: float, after: float, line: float, bold: bool):
    style.font.name = font
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    style.paragraph_format.keep_with_next = True


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    configure_style(normal, FONT_CJK, 11, INK, 0, 6, 1.25, False)
    configure_style(doc.styles["Heading 1"], FONT_CJK, 16, BLUE, 18, 10, 1.0, True)
    configure_style(doc.styles["Heading 2"], FONT_CJK, 13, BLUE, 14, 7, 1.0, True)
    configure_style(doc.styles["Heading 3"], FONT_CJK, 12, DARK_BLUE, 10, 5, 1.0, True)


def add_cover(doc: Document, header_lines: list[str], node_count: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(104)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V1.0  /  STORY MASTER")
    set_run_font(run, FONT_MONO, 10, TEAL, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第七次重启")
    set_run_font(run, FONT_CJK, 30, INK, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("8月4日人工重力整合版")
    set_run_font(run, FONT_CJK, 15, DARK_BLUE, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{node_count} NODES  /  EDITABLE REVIEW COPY")
    set_run_font(run, FONT_MONO, 9, GOLD, True)

    info = doc.add_paragraph()
    info.paragraph_format.left_indent = Inches(0.6)
    info.paragraph_format.right_indent = Inches(0.6)
    info.paragraph_format.space_before = Pt(0)
    info.paragraph_format.space_after = Pt(30)
    info.paragraph_format.line_spacing = 1.2
    shade_paragraph(info, SOFT_GRAY)
    set_paragraph_left_bar(info, "2E74B5", "16")
    visible_header = [line for line in header_lines if line and not line.startswith("#")]
    for index, line in enumerate(visible_header):
        if index:
            info.add_run("\n")
        run = info.add_run(line)
        set_run_font(run, FONT_CJK, 9, MUTED, bold=index < 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成于 {datetime.now().strftime('%Y-%m-%d %H:%M')}  /  来源：当前运行时剧情真源")
    set_run_font(run, FONT_CJK, 8.5, MUTED)
    doc.add_page_break()


def add_heading(doc: Document, line: str):
    if line.startswith("# "):
        text = line[2:].strip()
        style = "Heading 1"
    elif line.startswith("## "):
        text = line[3:].strip()
        style = "Heading 1"
    else:
        text = line.lstrip("#").strip()
        style = "Heading 2"
    paragraph = doc.add_paragraph(text, style=style)
    if style == "Heading 1":
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def add_node_block(doc: Document, block: list[str]):
    while block and block[0].startswith("#"):
        add_heading(doc, block.pop(0))
    if not block:
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.widow_control = True

    first = block[0]
    header = re.match(r"^\[([^\]]+)\]\s+\(([^/]+)/([^\)]+)\)$", first)
    start = 0
    if header:
        node_id, speaker, node_type = header.groups()
        run = paragraph.add_run(f"[{node_id}]  ")
        set_run_font(run, FONT_MONO, 8.4, TEAL, True)
        run = paragraph.add_run(f"{speaker} / {node_type}")
        set_run_font(run, FONT_CJK, 8.4, MUTED, True)
        start = 1

    for line in block[start:]:
        paragraph.add_run("\n")
        stripped = line.strip()
        if stripped.startswith("meta:"):
            run = paragraph.add_run(stripped)
            set_run_font(run, FONT_MONO, 7.8, MUTED)
        elif stripped.startswith("next:") or stripped.startswith("※"):
            run = paragraph.add_run(stripped)
            set_run_font(run, FONT_MONO, 7.9, MUTED)
        elif re.match(r"^\[[A-Z]\]\s*→", stripped):
            run = paragraph.add_run(stripped)
            set_run_font(run, FONT_CJK, 8.8, DARK_BLUE, True)
        elif stripped.startswith("文件标题:") or stripped.startswith("草稿标题:"):
            run = paragraph.add_run(stripped)
            set_run_font(run, FONT_CJK, 9.1, GOLD, True)
        else:
            run = paragraph.add_run(stripped)
            set_run_font(run, FONT_CJK, 9.2, INK)

    if header:
        shade_paragraph(paragraph, "F8FAFC")


def add_body(doc: Document, body_lines: list[str]):
    block: list[str] = []
    for line in body_lines + [""]:
        if line == "---":
            if block:
                add_node_block(doc, block)
                block = []
            doc.add_page_break()
            continue
        if line:
            block.append(line)
            continue
        if block:
            add_node_block(doc, block)
            block = []


def main():
    lines = SOURCE.read_text(encoding="utf-8-sig").splitlines()
    divider = lines.index("---")
    header_lines = lines[:divider]
    body_lines = lines[divider + 1:]
    node_ids = {
        match.group(1)
        for line in body_lines
        if (match := re.match(r"^\[([^\]]+)\]\s+\([^/]+/[^\)]+\)$", line))
    }
    node_ids.discard("MENU")
    node_count = len(node_ids)

    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "第七次重启 V1.0 主流程 8月4日人工重力整合版"
    doc.core_properties.subject = f"{node_count} 节点运行时剧情审校稿"
    doc.core_properties.author = "Seventh Reboot Project"
    doc.core_properties.keywords = "第七次重启, V1.0, 8月4日, 人工重力, Observer-01, Nova"

    add_cover(doc, header_lines, node_count)
    add_body(doc, body_lines)
    doc.save(OUTPUT)
    shutil.copy2(OUTPUT, DESKTOP_OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {DESKTOP_OUTPUT}")


if __name__ == "__main__":
    main()
